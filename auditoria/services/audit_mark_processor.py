"""
Procesador de Marcas de Auditoría - Inyecta marcas en documentos descargados

⚠️ PROTOCOLO DE SEGURIDAD CRÍTICO:
- SOLO modifica section.footer (documentos Word)
- SOLO agrega nuevas filas DESPUÉS de max_row (documentos Excel)
- NUNCA toca doc.paragraphs, doc.tables, section.header
- Devuelve documento sin cambios si ocurre algún error
"""

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl.styles import Font, PatternFill, Alignment
from django.db.models import Q
from auditoria.models import AuditMark
import re
import logging

logger = logging.getLogger(__name__)


class AuditMarkProcessor:
    """
    Procesa documentos para inyectar marcas de auditoría coincidentes.

    ⚠️ REGLAS DE SEGURIDAD CRÍTICAS:
    1. NUNCA modificar cuerpo del documento, tablas o encabezados
    2. SOLO agregar/modificar contenido de pie de página
    3. Siempre devolver documento sin cambios si ocurre algún error
    4. Verificar casos especiales ANTES del procesamiento
    """

    def __init__(self, audit_id, filename):
        self.audit_id = audit_id
        self.filename = filename
        self.normalized_filename = self.normalize_text(filename)

    @staticmethod
    def normalize_text(text):
        """
        Normalizar texto para emparejamiento.

        Ejemplos:
            'A 1 Balance.docx' → 'A1BALANCEDOCX'
            'A-1' → 'A1'
            '10 INTEGRACION EGRESOS' → '10INTEGRACIONEGRESOS'

        Args:
            text: Texto a normalizar

        Returns:
            str: Texto normalizado (mayúsculas, solo alfanumérico)
        """
        if not text:
            return ''
        text = text.upper()
        text = re.sub(r'[^A-Z0-9]', '', text)
        return text

    def get_matching_marks(self):
        """
        Obtener marcas que coincidan con el nombre de archivo de este documento.

        ⚠️ SEGURIDAD: Excluye marcas "Ejemplo:" a nivel de consulta

        Returns:
            list[AuditMark]: Lista de marcas coincidentes
        """
        # Consultar marcas activas, excluir ejemplos
        marks = AuditMark.objects.filter(
            audit_id=self.audit_id,
            is_active=True
        ).exclude(
            Q(description__icontains='Ejemplo:') |
            Q(description__icontains='Example:')
        )

        # Filtrar por emparejamiento de nombre de archivo
        matched_marks = []
        for mark in marks:
            if not mark.work_paper_number:
                continue

            normalized_wp = self.normalize_text(mark.work_paper_number)

            # Emparejamiento de subcadena bidireccional
            if (normalized_wp in self.normalized_filename or
                self.normalized_filename in normalized_wp):
                matched_marks.append(mark)
                logger.debug(
                    f"COINCIDENCIA: {mark.work_paper_number} ↔ {self.filename}"
                )

        return matched_marks

    def process_word_document(self, doc):
        """
        Agregar marcas de auditoría al pie de página del documento Word.

        ⚠️ VERIFICACIONES DE SEGURIDAD CRÍTICAS:
        1. Verificar si es archivo PROGRAMA → devolver sin cambios
        2. Obtener marcas coincidentes
        3. Verificar si hay coincidencias → si no, devolver sin cambios
        4. Solo modificar section.footer (nunca cuerpo/tablas/encabezados)
        5. Si ocurre algún error → devolver sin cambios

        Args:
            doc: Objeto Document de python-docx

        Returns:
            Document: Documento modificado (o sin cambios si hay errores/sin coincidencias)
        """
        try:
            # ═══════════════════════════════════════════════════════
            # CASO ESPECIAL 1: Archivos PROGRAMA
            # ═══════════════════════════════════════════════════════
            if 'PROGRAMA' in self.filename.upper():
                logger.info(
                    f"Archivo PROGRAMA detectado: {self.filename} "
                    f"- Omitiendo procesamiento de marcas para preservar hipervínculos"
                )
                return doc  # ⚠️ Devolver SIN CAMBIOS

            # ═══════════════════════════════════════════════════════
            # Obtener marcas coincidentes (excluye "Ejemplo:" a nivel de consulta)
            # ═══════════════════════════════════════════════════════
            matched_marks = self.get_matching_marks()

            # ═══════════════════════════════════════════════════════
            # CASO ESPECIAL 2: No se encontraron coincidencias
            # ═══════════════════════════════════════════════════════
            if not matched_marks or len(matched_marks) == 0:
                logger.info(
                    f"No se encontraron marcas coincidentes para {self.filename} "
                    f"- Omitiendo pie de página para evitar pie de página vacío"
                )
                return doc  # ⚠️ Devolver SIN CAMBIOS

            logger.info(
                f"Se encontraron {len(matched_marks)} marcas coincidentes para {self.filename}"
            )

            # ═══════════════════════════════════════════════════════
            # OPERACIÓN SEGURA: Agregar marcas SOLO al pie de página
            # ═══════════════════════════════════════════════════════
            self._add_marks_to_footer(doc, matched_marks)

            return doc

        except Exception as e:
            # ⚠️ CRÍTICO: Si ocurre CUALQUIER error, devolver documento sin cambios
            logger.error(
                f"Error al procesar marcas de auditoría para {self.filename}: {e}",
                exc_info=True
            )
            return doc  # ⚠️ Devolver SIN CAMBIOS

    def _add_marks_to_footer(self, doc, marks):
        """
        Agregar marcas al pie de página del documento.

        ⚠️ SEGURIDAD: SOLO modifica section.footer
        NUNCA toca doc.paragraphs, doc.tables, section.header

        Args:
            doc: Objeto Document
            marks: Lista de objetos AuditMark
        """
        # Agregar al pie de página en TODAS las secciones
        for section_idx, section in enumerate(doc.sections):
            try:
                footer = section.footer

                # Agregar línea en blanco para espaciado
                footer.add_paragraph()

                # Agregar párrafo de título
                p_title = footer.add_paragraph()
                run_title = p_title.add_run('MARCAS DE AUDITORÍA UTILIZADAS:')
                run_title.bold = True
                run_title.font.size = Pt(11)
                run_title.font.color.rgb = RGBColor(0, 112, 192)  # Azul profesional

                # Agregar cada marca como párrafo separado
                for mark in marks:
                    p_mark = footer.add_paragraph()
                    mark_text = f'{mark.symbol}  {mark.description}'
                    run_mark = p_mark.add_run(mark_text)
                    run_mark.bold = True
                    run_mark.font.size = Pt(11)
                    run_mark.font.color.rgb = RGBColor(0, 112, 192)

                logger.debug(
                    f"Agregadas {len(marks)} marcas al pie de página de la sección {section_idx + 1}"
                )

            except Exception as e:
                # ⚠️ Si la modificación del pie de página falla, registrar pero continuar
                # (el documento aún es utilizable sin marcas)
                logger.warning(
                    f"Error al agregar marcas a la sección {section_idx + 1}: {e}"
                )
                continue

    def process_excel_document(self, wb):
        """
        Agregar marcas de auditoría al final de la hoja Excel.

        ⚠️ VERIFICACIONES DE SEGURIDAD CRÍTICAS:
        1. Verificar si es archivo PROGRAMA → devolver sin cambios
        2. Obtener marcas coincidentes
        3. Verificar si hay coincidencias → si no, devolver sin cambios
        4. SOLO agregar nuevas filas DESPUÉS de max_row (nunca modificar existentes)
        5. Si ocurre algún error → devolver sin cambios

        Args:
            wb: Objeto Workbook de openpyxl

        Returns:
            Workbook: Libro de trabajo modificado (o sin cambios si hay errores/sin coincidencias)
        """
        try:
            # ═══════════════════════════════════════════════════════
            # CASO ESPECIAL 1: Archivos PROGRAMA
            # ═══════════════════════════════════════════════════════
            if 'PROGRAMA' in self.filename.upper():
                logger.info(
                    f"Archivo PROGRAMA detectado: {self.filename} "
                    f"- Omitiendo procesamiento de marcas"
                )
                return wb  # ⚠️ Devolver SIN CAMBIOS

            # ═══════════════════════════════════════════════════════
            # Obtener marcas coincidentes (excluye "Ejemplo:" a nivel de consulta)
            # ═══════════════════════════════════════════════════════
            matched_marks = self.get_matching_marks()

            # ═══════════════════════════════════════════════════════
            # CASO ESPECIAL 2: No se encontraron coincidencias
            # ═══════════════════════════════════════════════════════
            if not matched_marks or len(matched_marks) == 0:
                logger.info(
                    f"No se encontraron marcas coincidentes para {self.filename} "
                    f"- Omitiendo sección de marcas"
                )
                return wb  # ⚠️ Devolver SIN CAMBIOS

            logger.info(
                f"Se encontraron {len(matched_marks)} marcas coincidentes para {self.filename}"
            )

            # ═══════════════════════════════════════════════════════
            # OPERACIÓN SEGURA: Agregar marcas a NUEVAS FILAS SOLO
            # ═══════════════════════════════════════════════════════
            self._add_marks_to_excel_bottom(wb, matched_marks)

            return wb

        except Exception as e:
            # ⚠️ CRÍTICO: Si ocurre CUALQUIER error, devolver libro de trabajo sin cambios
            logger.error(
                f"Error al procesar marcas de auditoría para {self.filename}: {e}",
                exc_info=True
            )
            return wb  # ⚠️ Devolver SIN CAMBIOS

    def _add_marks_to_excel_bottom(self, wb, marks):
        """
        Agregar sección de marcas al final de la hoja Excel.

        ⚠️ SEGURIDAD: SOLO crea NUEVAS FILAS después de los datos existentes
        NUNCA modifica celdas existentes (fila <= max_row)

        Args:
            wb: Objeto Workbook
            marks: Lista de objetos AuditMark
        """
        ws = wb.active

        # ⚠️ CRÍTICO: Encontrar última fila con datos
        # Solo agregaremos filas DESPUÉS de esto
        last_row = ws.max_row

        # Comenzar nueva sección 3 filas después de los últimos datos
        start_row = last_row + 3

        logger.debug(
            f"Última fila de datos: {last_row}, "
            f"Iniciando sección de marcas en fila: {start_row}"
        )

        try:
            # Agregar fila de título (celdas combinadas B a E)
            ws.merge_cells(
                start_row=start_row,
                start_column=2,
                end_row=start_row,
                end_column=5
            )

            title_cell = ws.cell(row=start_row, column=2)
            title_cell.value = 'MARCAS DE AUDITORÍA UTILIZADAS:'
            title_cell.font = Font(bold=True, color='0070C0', size=11)
            title_cell.fill = PatternFill(
                start_color='D3D3D3',
                end_color='D3D3D3',
                fill_type='solid'
            )
            title_cell.alignment = Alignment(
                horizontal='left',
                vertical='center'
            )

            # Agregar cada marca como una fila
            for i, mark in enumerate(marks, start=1):
                row_num = start_row + i

                # ⚠️ VERIFICACIÓN DE SEGURIDAD: Asegurarse de que estamos creando una NUEVA fila
                if row_num <= last_row:
                    logger.error(
                        f"VIOLACIÓN DE SEGURIDAD: Se intentó modificar la fila existente {row_num}"
                    )
                    raise ValueError("No se pueden modificar filas existentes")

                # Crear nueva celda (seguro porque row_num > last_row)
                cell = ws.cell(row=row_num, column=2)
                cell.value = f'{mark.symbol}  {mark.description}'
                cell.font = Font(bold=True, color='0070C0', size=11)
                cell.fill = PatternFill(
                    start_color='E7E6E6',
                    end_color='E7E6E6',
                    fill_type='solid'
                )
                cell.alignment = Alignment(horizontal='left', vertical='center')

            logger.info(
                f"Agregadas {len(marks)} marcas a Excel comenzando en fila {start_row}"
            )

        except Exception as e:
            logger.error(f"Error al agregar marcas a Excel: {e}")
            raise  # Re-lanzar para activar try/except externo
